from   django.contrib                 import messages
from   django.views.decorators.csrf   import csrf_exempt
import json
from   django.http                    import HttpResponseBadRequest, JsonResponse
from   django.template.loader         import render_to_string  # Para usar templates si es necesario
from   django.utils.html              import strip_tags  # Para generar una versión de solo texto
from   datetime                       import datetime, timedelta
from   django.utils                   import timezone
from   django.core.mail               import send_mail
from   django.core.paginator          import Paginator
from   django.db.models               import Q
from   django.views.decorators.http   import require_POST
from   django.utils.timezone          import now
from   django.templatetags.static     import static
from   django.conf                    import settings
from   docx                           import Document
from   django.utils.timezone          import make_aware
from   docx.shared                    import Pt, RGBColor, Inches
from   docx.enum.text                 import WD_ALIGN_PARAGRAPH
from   docx.enum.table                import WD_TABLE_ALIGNMENT
from   docx.oxml                      import OxmlElement
from   docx.oxml.ns                   import qn
import os
import io
import pandas as pd
from   django.shortcuts               import get_object_or_404, render, redirect
from   django.contrib.auth.decorators import login_required
from   .forms                         import CSVUploadForm , ExcelUploadFrom, TareaForm
from   openpyxl.styles                import PatternFill
from   django.http                    import HttpResponse
from   datetime                       import date
from   calendar                       import monthrange
from   .models                        import Tarea, User, ControlActualizacionMensual, Cliente


@login_required
def home(request):
    return render(request, 'csv_processor/home.html')

def obtener_usuario_predeterminado():
    return User.objects.get(username='Eliana')  # o el username que hayas creado

    # en tu vista:
    usuario = request.user if request.user.is_authenticated else obtener_usuario_predeterminado()

@login_required
def procesar_excel(request):
    context = {}
    context['error'] = ''
    if request.method == "POST":
        form = CSVUploadForm (request.POST, request.FILES)
        if form.is_valid():
            # Obtener el formato del archivo seleccionado
            file_format = form.cleaned_data['file_format']
            archivo = request.FILES['csv_file']  # 'csv_file' será el archivo subido por el usuario

            # Validar que el archivo tenga la extensión '.xlsx' para Excel
            file_name, file_extension = os.path.splitext(archivo.name)

            if not validar_extension(file_extension, 1):
                context['form'] = form
                context['error'] = "Por favor, sube un archivo con extensión .xlsx o .xls"
                return render(request, 'csv_processor/procesar_excel.html', context)
            

            if file_format == 0:
                context['form'] = form
                context['error'] = "Por favor, seleccione un formato"
                return render(request, 'csv_processor/procesar_excel.html', context)
            
            pd.set_option('display.max_rows', None)

            # Obtener el directorio donde está el archivo CSV que se sube
            csv_directory = os.path.dirname(archivo.name)  # obtiene el directorio del archivo CSV

             # Crear la ruta completa para el archivo de salida en el mismo directorio
            output_file_name = os.path.join(csv_directory, archivo.name)  # Usa el mismo nombre de archivo
            
            # Leer el archivo Excel
            try:
                # Se usa pd.read_excel para leer archivos Excel
                df = pd.read_excel(archivo, engine='openpyxl')  # Puedes especificar el 'engine' si es necesario
                
                # Evitar que ponga decimales
                df = df.apply(pd.to_numeric, errors='ignore').astype('Int64', errors='ignore')

                # Reemplaza las celdas vacías por texto vacío ('') en todo el DataFrame
                df = df.fillna(0)
                df = df.applymap(lambda x: '' if x == 0 else x)

                df.columns = df.columns.str.strip()
                print(df.columns)
                # Procesar según el formato seleccionado
                if file_format == '1006':
                    # Convierte a numérico los valores de la columna impuestos descontables
                    df['Impuesto generado'] = pd.to_numeric(df['Impuesto generado'], errors='coerce')
                    
                    df['IVA recuperado en devoluciones en compras anuladas. rescindidas o resueltas'] = pd.to_numeric(df['IVA recuperado en devoluciones en compras anuladas. rescindidas o resueltas'], errors='coerce')
                    
                    df['Impuesto al consumo'] = pd.to_numeric(df['Impuesto al consumo'], errors='coerce')
                    
                    # Agrupa y suma los valores de la columna I
                    df_grouped_I = df.groupby(['Tipo de Documento', 'Número identificación', 'DV',
                                               'Primer apellido del informado', 'Segundo apellido del informado',
                                               'Primer nombre del informado', 'Otros nombres del informado',
                                               'Razón social informado'])['Impuesto generado'].sum().reset_index()
                    
                    # Agrupa y suma los valores de la columna J
                    df_grouped_J = df.groupby(['Tipo de Documento', 'Número identificación', 'DV',
                                               'Primer apellido del informado', 'Segundo apellido del informado',
                                               'Primer nombre del informado', 'Otros nombres del informado',
                                               'Razón social informado'])['IVA recuperado en devoluciones en compras anuladas. rescindidas o resueltas'].sum().reset_index()
                    
                    # Agrupa y suma los valores de la columna J
                    df_grouped_K = df.groupby(['Tipo de Documento', 'Número identificación', 'DV',
                                               'Primer apellido del informado', 'Segundo apellido del informado',
                                               'Primer nombre del informado', 'Otros nombres del informado',
                                               'Razón social informado'])['Impuesto al consumo'].sum().reset_index()

                    # Hace merge de las columnas sumadas con las demás columnas
                    merged_ij = pd.merge(df_grouped_I, df_grouped_J, on=[ 'Tipo de Documento', 'Número identificación', 'DV',
                                                                          'Primer apellido del informado', 'Segundo apellido del informado',
                                                                          'Primer nombre del informado', 'Otros nombres del informado',
                                                                          'Razón social informado'
                                                                        ], how='left')

                    df_grouped = pd.merge(merged_ij, df_grouped_K, on=[
                                                                        'Tipo de Documento', 'Número identificación', 'DV',
                                                                        'Primer apellido del informado', 'Segundo apellido del informado',
                                                                        'Primer nombre del informado', 'Otros nombres del informado',
                                                                        'Razón social informado'
                                                                      ], how='left')
                    
                    # Crea el archivo Excel en memoria
                    return crear_archivo_excel_respuesta(df_grouped, output_file_name, file_format)

                elif file_format == '1005':
                    # Convierte a numérico los valores de la columna impuestos descontables
                    df['Impuesto descontable'] = pd.to_numeric(df['Impuesto descontable'], errors='coerce')
                    
                    df['IVA resultante por devoluciones en ventas anuladas, rescindidas o resueltas'] = pd.to_numeric(df['IVA resultante por devoluciones en ventas anuladas, rescindidas o resueltas'], errors='coerce')
                    
                    # Agrupa y suma los valores de la columna I
                    df_grouped_I = df.groupby(['Numero de identificación del informado', 'Tipo de Documento', 'DV', 
                                             'Primer apellido del informado', 'Segundo apellido del informado',
                                             'Razón social informado'])['Impuesto descontable'].sum().reset_index()
                    
                    # Agrupa y suma los valores de la columna J
                    df_grouped_J = df.groupby(['Numero de identificación del informado', 'Tipo de Documento', 'DV', 
                                             'Primer apellido del informado', 'Segundo apellido del informado',
                                             'Razón social informado'])['IVA resultante por devoluciones en ventas anuladas, rescindidas o resueltas'].sum().reset_index()
                    
                    # Hace merge de las columnas sumadas con las demás columnas
                    df_grouped = pd.merge(df_grouped_I, df_grouped_J, on=['Numero de identificación del informado', 'Tipo de Documento', 'DV', 
                                                                          'Primer apellido del informado', 'Segundo apellido del informado',
                                                                          'Razón social informado'], how='left') 
                    
                    # Crea el archivo Excel en memoria
                    return crear_archivo_excel_respuesta(df_grouped, output_file_name, file_format)

                elif file_format == '1007':
                    # Convierte a numérico los valores de la columna impuestos descontables
                    df['Ingresos brutos recibidos'] = pd.to_numeric(df['Ingresos brutos recibidos '], errors='coerce')
                    
                    df['Devoluciones, rebajas y descuentos'] = pd.to_numeric(df['Devoluciones, rebajas y descuentos'], errors='coerce')

                    # Agrupar por NIT y sumar valores
                    # Agrupa y suma los valores de la columna I
                    df_grouped_I = df.groupby(['Concepto', 'Tipo de documento', 'Número identificación del informado',
                                                'Primer apellido del informado', 'Segundo apellido del informado',
                                                'Primer nombre del informado', 'Otros nombres del informado',
                                                'Razón social informado', 'País de residencia o domicilio'])['Ingresos brutos recibidos'].sum().reset_index()
                    
                    # Agrupa y suma los valores de la columna J
                    df_grouped_J = df.groupby(['Concepto', 'Tipo de documento', 'Número identificación del informado',
                                                'Primer apellido del informado', 'Segundo apellido del informado',
                                                'Primer nombre del informado', 'Otros nombres del informado',
                                                'Razón social informado', 'País de residencia o domicilio'])['Devoluciones, rebajas y descuentos'].sum().reset_index()
                     
                     # Hace merge de las columnas sumadas con las demás columnas
                    df_grouped = pd.merge(df_grouped_I, df_grouped_J, on=['Concepto', 'Tipo de documento', 'Número identificación del informado',
                                                                           'Primer apellido del informado', 'Segundo apellido del informado',
                                                                           'Primer nombre del informado', 'Otros nombres del informado',
                                                                           'Razón social informado', 'País de residencia o domicilio'], how='left') 
                     
                     # Crea el archivo Excel en memoria
                    return crear_archivo_excel_respuesta(df_grouped, output_file_name, file_format)

            except Exception as e:
                context['error'] = f"Ocurrió un error al procesar el archivo: {e}"

    else:
        form = CSVUploadForm ()

    context['form'] = form
    
    return render(request, 'csv_processor/procesar_excel.html', context)

@login_required
def proveedores(request):
    context = {}  # Definimos context al principio

    if request.method == 'POST':
        
        form = ExcelUploadFrom(request.POST, request.FILES)

        archivo = request.FILES.get('excel_file_proveedor')  # Obtenemos el archivo cargado
        if not archivo:
            context['form'] = form
            context['error'] = "Por favor, seleccione un archivo"
            return render(request, 'csv_processor/proveedores.html', context)

        # Intentar leer las hojas del archivo Excel
        try:
            df_1001 = pd.read_excel(archivo, sheet_name="1001")
            df_terceros = pd.read_excel(archivo, sheet_name="terceros proveedores")
        except ValueError:
            context['form'] = form
            context['error'] = "El archivo no contiene las hojas requeridas: '1001' y 'terceros_proveedores'."
            return render(request, 'csv_processor/proveedores.html', context)

        # Verificar si las columnas necesarias existen en ambas hojas
        if 'Número identificación del informado' not in df_1001.columns or 'nit_ter' not in df_terceros.columns:
            context['form'] = form
            context['error'] = "Las columnas 'Número identificación del informado' o 'nit_ter' no existen en el archivo"
            return render(request, 'csv_processor/proveedores.html', context)

        
        # Limpiar la columna 'nit_ter' quitando las comas y eliminando espacios antes de la validación
        df_terceros['nit_ter'] = df_terceros['nit_ter'].astype(str).str.replace(',', '', regex=False).str.strip()

        # Limpiar la columna 'Número identificación del informado' eliminando espacios
        df_1001['Número identificación del informado'] = df_1001['Número identificación del informado'].astype(str).str.strip()

        # Validación y marcado de proveedores
        proveedores_validos = df_1001['Número identificación del informado'].isin(df_terceros['nit_ter'])

        # Crear un archivo Excel con el resultado
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename="resultado_proveedores.xlsx"'

        # Crear un workbook de openpyxl
        with pd.ExcelWriter(response, engine='openpyxl') as writer:
            df_1001.to_excel(writer, sheet_name='1001', index=False)
            df_terceros.to_excel(writer, sheet_name='terceros_proveedores', index=False)
            
            # Acceder al archivo Excel ya creado
            workbook = writer.book
            worksheet_1001 = workbook['1001']

            # Establecer el color azul suave para los proveedores
            azul_suave = PatternFill(start_color="BCE0F5", end_color="BCE0F5", fill_type="solid")

            # Colorear las filas de los proveedores en azul suave
            for row in range(2, len(df_1001) + 2):  # Comenzamos en 2 para evitar la fila de cabeceras
                if proveedores_validos.iloc[row - 2]:  # Si la fila es un proveedor
                    for col in range(1, len(df_1001.columns) + 1):  # Recorremos todas las columnas
                        worksheet_1001.cell(row=row, column=col).fill = azul_suave

        return response

    else:
        # Si no es un POST, renderizamos el formulario vacío
        form = ExcelUploadFrom()
        context['form'] = form

    return render(request, 'csv_processor/proveedores.html', context)


#Funcion para validar la extensión del archivo
def validar_extension(extension, flag):
    if flag == 0:
        return extension.lower() == '.csv'
    elif flag == 1:
        return extension.lower() == '.xlsx'

#lee el archio
def leer_archivo(archivo, flagTipoArchivo):
    
    """Lee el archivo CSV y lo convierte en un DataFrame de pandas."""
    if flagTipoArchivo == 0:
        try:
            return pd.read_csv(archivo, encoding='utf-8', delimiter=';')
        except UnicodeDecodeError:
            try:
                return pd.read_csv(archivo, encoding='ISO-8859-1', delimiter=';')
            except UnicodeDecodeError:
                return pd.read_csv(archivo, encoding='Windows-1252', delimiter=';')
    elif flagTipoArchivo == 1:
         """Lee el archivo Excel y lo convierte en un DataFrame de pandas."""
         return pd.read_excel(archivo)
"""
Crea un archivo Excel en memoria a partir de un DataFrame y devuelve una respuesta HTTP 
que permite descargar el archivo Excel generado."""

def crear_archivo_excel_respuesta(df, output_file_name, file_sheet):
   
    """
    :param df: El DataFrame de pandas que contiene los datos a cdincluir en el archivo Excel.
    :param output_file_name: El nombre del archivo Excel que se enviará como descarga.
    :return: Respuesta HTTP con el archivo Excel generado.
    """
    # Crea el archivo Excel en memoria
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name= file_sheet)  # Escribe el DataFrame en la hoja 'Resumen'

    # Prepara la respuesta HTTP para enviar el archivo como una descarga
    response = HttpResponse(output.getvalue(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="{output_file_name}"'
    return response

#Vista para gestionar tareas
@login_required
def tablero_kanban(request):
    try:
        hoy        = date.today()
        anio       = int(request.GET.get('anio', hoy.year))
        mes        = int(request.GET.get('mes', hoy.month))
        ahora      = timezone.now()

        primer_dia = date(anio, mes, 1)
        ultimo_dia = primer_dia.replace(day=monthrange(anio, mes)[1])
        usuario = request.user

        # Buscar último registro de actualización
        registro, creado = ControlActualizacionMensual.objects.get_or_create(
            usuario=usuario,
            defaults={'ultima_actualizacion': hoy}
        )
        necesita_actualizacion = False

        # Primera vez o no actualizado este mes
        if creado or registro.ultima_actualizacion < primer_dia:
            # Si es la primera vez (sin importar el día) o si es día 1 del mes
            if creado or hoy.day == 1:
                necesita_actualizacion = True
                primer_dia_mes_anterior = (primer_dia - timedelta(days=1)).replace(day=1)
                ultimo_dia_mes_anterior = primer_dia - timedelta(days=1)

                tareas_a_mover = Tarea.objects.filter(
                    fecha__range=(primer_dia_mes_anterior, ultimo_dia_mes_anterior),
                    estado__in=['pendiente', 'en_progreso'],
                    usuario=usuario
                )
                tareas_a_mover.update(fecha=primer_dia)

                # Actualizar el registro
                registro.ultima_actualizacion = hoy
                registro.save()

        # Consultas de tareas
        tareas_pendientes = Tarea.objects.filter(
            fecha__range=(primer_dia, ultimo_dia),
            estado='pendiente',
            usuario=usuario
        ).order_by('fecha_vencimiento')

        tareas_en_progreso = Tarea.objects.filter(
            fecha__range=(primer_dia, ultimo_dia),
            estado='en_progreso',
            usuario=usuario
        ).order_by('fecha_vencimiento')

        tareas_completadas = Tarea.objects.filter(
            Q(
                fecha__range=(primer_dia, ultimo_dia),
                estado='completada',
                usuario=usuario
            ) & (
                Q(fecha_completado__isnull=True) |
                Q(fecha_completado__gt=ahora - timedelta(days=1))
            )
        ).order_by('fecha_vencimiento')

        clientes = Cliente.objects.filter(contador=usuario)

        context = {
            'hoy': primer_dia,
            'anio': anio,
            'mes': mes,
            'tareas_por_estado': {
                'pendiente': tareas_pendientes,
                'en_progreso': tareas_en_progreso,
                'completada': tareas_completadas,
            },
            'clientes': clientes,
        }
        return render(request, 'csv_processor/kanban.html', context)

    except Exception as e:
        import traceback
        print("⚠️ Error en vista kanban:", e)
        traceback.print_exc()
        return HttpResponse("Error interno del servidor", status=500)
    
@login_required
def crear_tarea(request):
    clientes = Cliente.objects.filter(contador=request.user) if request.user.is_authenticated else []

    if request.method == 'POST':
        nombre       = request.POST.get('nombre')
        descripcion  = request.POST.get('descripcion')
        fecha        = request.POST.get('fecha')
        fechavence   = request.POST.get('fecha_vencimiento')
        cliente_id   = request.POST.get('cliente')

        if nombre and descripcion and fecha:
            try:
                fecha_creacion     = date.fromisoformat(fecha)
                fecha_vencimiento  = date.fromisoformat(fechavence)

                if fecha_vencimiento < fecha_creacion or fecha_vencimiento < date.today():
                    error = "La fecha de vencimiento no puede ser menor a la fecha de creación ni a la fecha actual."
                    return render(request, 'csv_processor/crear_tarea.html', {
                        'error': error,
                        'nombre': nombre,
                        'descripcion': descripcion,
                        'fecha': fecha,
                        'fecha_vencimiento': fechavence,
                        'clientes': clientes,
                    })

                cliente = None
                if cliente_id:
                    try:
                        cliente = Cliente.objects.get(id=cliente_id, contador=request.user)
                    except Cliente.DoesNotExist:
                        error = "El cliente seleccionado no es válido."
                        return render(request, 'csv_processor/crear_tarea.html', {
                            'error': error,
                            'nombre': nombre,
                            'descripcion': descripcion,
                            'fecha': fecha,
                            'fecha_vencimiento': fechavence,
                            'clientes': clientes,
                        })

                if request.user.is_authenticated:
                    usuario = request.user

                    tarea = Tarea(
                        titulo            = nombre,
                        descripcion       = descripcion,
                        fecha             = fecha_creacion,
                        fecha_vencimiento = fecha_vencimiento,
                        usuario           = usuario,
                        cliente           = cliente
                    )
                    tarea.save()
                    return redirect('kanban')

            except Exception as e:
                print("Error al guardar la tarea:", e)
        else:
            print("Datos incompletos")

    return render(request, 'csv_processor/crear_tarea.html', {
        'clientes': clientes,
    })

@csrf_exempt
def actualizar_estado_tarea(request):
    if request.method == 'POST':
        try:
            data         = json.loads(request.body)
            tarea_id     = data.get('tarea_id') or data.get('id')
            nuevo_estado = data.get('estado')
            tarea        = Tarea.objects.get(id=tarea_id)

            if tarea.estado != 'completada' and nuevo_estado == 'completada':
                tarea.fecha_completado = timezone.now()
            elif nuevo_estado != 'completada':
                tarea.fecha_completado = None  # Resetear si vuelve a otro estado

            tarea.estado = nuevo_estado
            print(tarea.estado)
            tarea.save()
            
            return JsonResponse({'success': True})
        except Exception as e:
            print("Error al actualizar la tarea:", str(e))
            return JsonResponse({'success': False, 'error': str(e)})

    return JsonResponse({'success': False, 'error': 'Método no permitido'}, status=405)


@login_required
@csrf_exempt  # ⚠️ Temporal, mientras implementás CSRF bien con fetch
def editar_tarea(request, tarea_id):
    tarea = get_object_or_404(Tarea, id=tarea_id)

    if request.method == 'POST':
        try:
            data                    = json.loads(request.body)
            tarea.titulo            = data.get('titulo')
            tarea.descripcion       = data.get('descripcion')
            tarea.fecha_vencimiento = data.get('fecha')

            cliente_id = data.get('cliente_id')
            if cliente_id:
                cliente = get_object_or_404(Cliente, id=cliente_id)
                tarea.cliente = cliente
            else:
                tarea.cliente = None  # Si viene vacío, se desasocia

            tarea.save()
            return JsonResponse({'success': True})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    
    return JsonResponse({'error': 'Método no permitido'}, status=405)

@login_required
@csrf_exempt
def eliminar_tarea(request, tarea_id):
    try:
        if request.method == "DELETE":
            # Obtener la tarea que queremos eliminar
            tarea = get_object_or_404(Tarea, id=tarea_id)
            # Eliminar la tarea de la base de datos
            tarea.delete()
        
            return JsonResponse({'success': True})
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})
    
    return JsonResponse({'error': 'Método no permitido'}, status=405)

@login_required
def historial_tareas_completadas(request):
    usuario      = request.user
    clientes = Cliente.objects.filter(contador=usuario).order_by('nombre')
    tareas_lista = Tarea.objects.filter(
        estado   = 'completada',
        usuario  = usuario
    ).order_by('-fecha_completado')

    paginator   = Paginator(tareas_lista, 10)  # 10 tareas por página
    page_number = request.GET.get('page')
    page_obj    = paginator.get_page(page_number)

    context = {
        'page_obj': page_obj,
        'clientes': clientes,
    }
    return render(request, 'csv_processor/historial_tareas.html', context)

@require_POST
@login_required
def cambiar_estado_tarea(request, tarea_id):
    nuevo_estado = request.POST.get('estado')
    tarea        = get_object_or_404(Tarea, id=tarea_id, usuario=request.user)

    if nuevo_estado in ['pendiente', 'en_progreso']:
        tarea.estado           = nuevo_estado
        tarea.fecha_completado = None  # quitamos fecha de completado si se revierte
        tarea.save()
        return redirect('historial_tareas')
    
    return HttpResponseBadRequest("Estado inválido")

#Envia correos con las tareas pendientes o en proceso
@csrf_exempt
def enviar_tareas(request):
    
    usuarios = User.objects.all()
    
    # Obtener la fecha actual
    fecha_actual      = datetime.now().date()
    # Definir el rango de "próxima vencimiento" (por ejemplo, 15 días)
    rango_vencimiento = fecha_actual + timedelta(days=15)
    
    for usuario in usuarios:
        tareas = Tarea.objects.filter(
            estado__in=['pendiente', 'en_progreso'],
            usuario=usuario
        ).order_by('fecha_vencimiento')  # Ordenar por fecha de vencimiento ascendente
        
        if tareas.exists():
            # Construir la tabla HTML con las tareas
            tabla_html = """
            <table border="1" cellpadding="10" cellspacing="0" style="border-collapse: collapse; width: 100%; margin-top: 20px;">
                <thead>
                    <tr>
                        <th>Título</th>
                        <th>Descripción</th>
                        <th>Fecha de Creación</th>
                        <th>Fecha de Vencimiento</th>
                    </tr>
                </thead>
                <tbody>
            """
            
            for t in tareas:
                fila_color = 'background-color: #FFCCCB;' if t.fecha_vencimiento and t.fecha_vencimiento <= rango_vencimiento else ''
                tabla_html += f"""
                <tr style="{fila_color}">
                    <td>{t.titulo}</td>
                    <td>{t.descripcion}</td>
                    <td>{t.fecha}</td>
                    <td>{t.fecha_vencimiento if t.fecha_vencimiento else 'No tiene'}</td>
                </tr>
                """
            
            tabla_html += """
                </tbody>
            </table>
            """

            # Footer del correo
            footer_html = """
            <hr>
            <p style="font-size: 12px; color: gray;">
                Este correo fue enviado automáticamente por la plataforma de productividad contable Accountants Tools.<br>
                Si tienes preguntas, contáctanos a <a href="mailto:fabricio.galarza@outlook.com">fabricio.galarza@outlook.com.com</a>.
            </p>
            """

            # Cuerpo completo del correo
            mensaje_html = f"""
            <p>Hola {usuario.first_name},</p>
            <p>Estas son las tareas que tienes pendientes o en progreso:</p>
            {tabla_html}
            <p>¡Saludos!</p>
            {footer_html}
            """

            send_mail(
                'TIENES TAREAS SIN FINALIZAR',
                strip_tags(mensaje_html),  # Versión de solo texto
                'fabricio.galarzadev@gmail.com',
                [usuario.email],
                fail_silently=False,
                html_message=mensaje_html
            )

    return JsonResponse({'status': 'correos enviados'})

@login_required
def crear_cliente(request):
    if request.method == 'POST':
        nombre = request.POST.get('nombre')
        identificacion = request.POST.get('identificacion')
        email = request.POST.get('email')
        telefono = request.POST.get('telefono')
        direccion = request.POST.get('direccion')

        if Cliente.objects.filter(identificacion=identificacion, contador=request.user).exists():
            messages.error(request, "Ya existe un cliente con esa identificación.")
        else:
            Cliente.objects.create(
                nombre=nombre,
                identificacion=identificacion,
                email=email,
                telefono=telefono,
                direccion=direccion,
                contador=request.user
            )
            messages.success(request, "Cliente registrado exitosamente.")
            return redirect('crear_cliente')

    return render(request, 'csv_processor/crear_cliente.html')  # sin redefinir 'messages'

@login_required
def exportar_reporte_cliente(request):
    cliente_id = request.GET.get('cliente_id')
    estado = request.GET.get('estado')
    fecha_inicio = request.GET.get('fecha_inicio')
    fecha_fin = request.GET.get('fecha_fin')
    hoy = now().date()

    try:
        cliente = Cliente.objects.get(id=cliente_id)
    except Cliente.DoesNotExist:
        return HttpResponse("Cliente no encontrado", status=404)

    # Determinar fechas
    if fecha_inicio:
        primer_dia = make_aware(datetime.combine(datetime.strptime(fecha_inicio, '%Y-%m-%d'), datetime.min.time()))
    else:
        primer_dia = make_aware(datetime.combine(hoy.replace(day=1), datetime.min.time()))

    if fecha_fin:
        ultimo_dia = make_aware(datetime.combine(datetime.strptime(fecha_fin, '%Y-%m-%d'), datetime.max.time()))
    else:
        ultimo_dia = make_aware(datetime.combine(hoy, datetime.max.time()))

    # Construir filtro base
    filtro = {'cliente': cliente}

    if estado:
        filtro['estado'] = estado
        if estado == 'completado':
            filtro['fecha_completado__range'] = (primer_dia, ultimo_dia)
    else:
        filtro['fecha_completado__range'] = (primer_dia, ultimo_dia)

    print("Estado recibido:", estado)
    print("Filtro :", filtro)
    tareas = Tarea.objects.filter(**filtro).order_by('fecha_completado')

    doc = Document()

    # === Agregar tabla para logo y nombre del usuario logueado ===
    tabla_header = doc.add_table(rows=1, cols=2)
    tabla_header.alignment = WD_TABLE_ALIGNMENT.LEFT

    perfil = request.user.profile  # Asegúrate que el perfil exista (por señal o creación manual)

    profesion = perfil.profesion or 'Profesión no especificada'
    area = perfil.areaOperativa or 'Área no especificada'
    logo_nombre = perfil.nombreLogo or 'logo_default.png'  # nombre del archivo, por ejemplo
    # Colocar logo en la primera celda (izquierda)
    print(logo_nombre)
    logo_path = os.path.join(settings.BASE_DIR, 'static', 'images', logo_nombre)  # Asegúrate de tener este logo
    if os.path.exists(logo_path):
        cell_logo = tabla_header.cell(0, 0)
        cell_logo.width = Inches(1.0)  # Reducir el ancho de la celda para el logo
        cell_logo.paragraphs[0].add_run().add_picture(logo_path, width=Inches(1.0))  # Ajustamos el tamaño del logo
        cell_logo.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Colocar nombre del usuario logueado en la segunda celda (derecha)
    cell_nombre_usuario = tabla_header.cell(0, 1)
    nombre_usuario_parrafo = cell_nombre_usuario.paragraphs[0]
    
    # Obtener el nombre del usuario logueado (nombre completo o username)
    nombre_usuario = request.user.get_full_name() if request.user.get_full_name() else request.user.username
    
    # Agregar el nombre del usuario al lado derecho del logo
    run = nombre_usuario_parrafo.add_run({nombre_usuario})
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(204, 51, 51 )
    nombre_usuario_parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Alineación izquierda para evitar separación

    # Ajustar el ancho de la segunda celda para evitar separación
    cell_nombre_usuario.width = Inches(3.0)  # Ajustar el ancho de la celda para el nombre del usuario

    

    # Crear un nuevo párrafo debajo
    parrafo_profesion = cell_nombre_usuario.add_paragraph()
    run_profesion = parrafo_profesion.add_run({profesion})
    run_profesion.font.size = Pt(11)
    run_profesion.font.name = 'Arial'
    run_profesion.font.color.rgb = RGBColor(100, 100, 100)  # Gris oscuro
    run_profesion.italic = True  # Aplicar itálica
    parrafo_profesion.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # === Encabezado con tabla (mismos estilos previos) ===
    encabezado_tabla = doc.add_table(rows=1, cols=1)
    encabezado_tabla.alignment = WD_TABLE_ALIGNMENT.CENTER  # Centrar toda la tabla
    cell = encabezado_tabla.cell(0, 0)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    run = p.add_run(cliente.nombre.upper() + "\n")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = 'Arial'

    run2 = p.add_run("REPORTE DE ACTIVIDADES\n")
    run2.bold = True
    run2.font.size = Pt(14)
    run2.font.name = 'Arial'

    run3 = p.add_run(f"Período: {primer_dia.date()} a {ultimo_dia.date()}\n")
    run3.font.size = Pt(11)
    run3.font.name = 'Arial'

    run4 = p.add_run(f"Departamento: {area}")
    run4.font.size = Pt(11)
    run4.italic = True
    run4.font.name = 'Arial'

    # Agregar borde inferior a la celda
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')  # Un poco más grueso
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    borders.append(bottom)
    tc_pr.append(borders)

    # Espacio después del encabezado
    doc.add_paragraph("")

    if tareas.exists():
        tabla = doc.add_table(rows=1, cols=2)
        tabla.style = 'Table Grid'
        hdr_cells = tabla.rows[0].cells
        hdr_cells[0].text = 'Título'
        hdr_cells[1].text = 'Descripción'

        for tarea in tareas:
            row_cells = tabla.add_row().cells
            row_cells[0].text = tarea.titulo
            row_cells[1].text = tarea.descripcion
    else:
        doc.add_paragraph("No se encontraron tareas para este cliente en el período y estado seleccionados.")

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=Reporte_{cliente.nombre}.docx'
    doc.save(response)
    return response